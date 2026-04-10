import streamlit as st
from openai import OpenAI
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# ---------- NVIDIA API ----------
client = OpenAI(
    api_key=st.secrets["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
 "],
    base_url="https://integrate.api.nvidia.com/v1"
)

st.set_page_config(page_title="AI DRD Agent", layout="wide")
st.title("🤖 AI Defect Resolution Agent (Enterprise Level)")

# ---------- MEMORY ----------
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# ---------- INPUT ----------
user_input = st.text_area("Enter Defect / JIRA Details", height=150)

# ---------- SECTION EXTRACT ----------
def extract_section(text, section):
    lines = text.split("\n")
    capture = False
    result = []

    for line in lines:
        if section.lower() in line.lower():
            capture = True
            continue
        if capture:
            if line.strip() == "":
                break
            result.append(line)

    return "\n".join(result)

# ---------- CREATE DOCX ----------
def create_docx(output):
    doc = Document()
    doc.add_heading('Defect Resolution Document', 0)

    doc.add_heading('Description', 1)
    doc.add_paragraph(extract_section(output, "Description"))

    doc.add_heading('Analysis', 1)
    doc.add_paragraph(extract_section(output, "Analysis"))

    doc.add_heading('Proposed Solution', 1)
    doc.add_paragraph(extract_section(output, "Proposed Solution"))

    doc.add_heading('Testing Approach', 1)
    doc.add_paragraph(extract_section(output, "Testing"))

    doc.add_heading('Impact Analysis', 1)
    doc.add_paragraph(extract_section(output, "Impact"))

    doc.add_heading('Release Notes', 1)
    doc.add_paragraph(extract_section(output, "Release"))

    file_name = "DRD_Report.docx"
    doc.save(file_name)
    return file_name

# ---------- CREATE PDF ----------
def create_pdf(output):
    file_name = "DRD_Report.pdf"
    c = canvas.Canvas(file_name, pagesize=letter)

    text = c.beginText(40, 750)
    text.setFont("Helvetica", 10)

    for line in output.split("\n"):
        text.textLine(line)

    c.drawText(text)
    c.save()

    return file_name

# ---------- BUTTON ----------
if st.button("Run AI Agent"):

    st.session_state.chat_history.append({"role": "user", "content": user_input})

    messages = [
        {
            "role": "system",
            "content": """
You are an AI Defect Resolution Agent.

Steps:
1. Analyze defect
2. If missing info → ask questions
3. If enough → generate DRD

Format:
Description:
Analysis:
Proposed Solution:
Testing Approach:
Impact Analysis:
Release Notes:
"""
        }
    ] + st.session_state.chat_history

    response = client.chat.completions.create(
        model="meta/llama3-70b-instruct",
        messages=messages,
        temperature=0.3
    )

    reply = response.choices[0].message.content

    st.session_state.chat_history.append({"role": "assistant", "content": reply})

    st.subheader("🤖 Agent Output")
    st.write(reply)

    # Create files
    docx_file = create_docx(reply)
    pdf_file = create_pdf(reply)

    # Download buttons
    with open(docx_file, "rb") as f:
        st.download_button("📄 Download Word", f, file_name="DRD.docx")

    with open(pdf_file, "rb") as f:
        st.download_button("📄 Download PDF", f, file_name="DRD.pdf")